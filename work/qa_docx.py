from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "docx",
        nargs="?",
        default="outputs/Graph_Rescue_RAG_Theory.docx",
    )
    parser.add_argument(
        "--report",
        default="work/docx_structural_qa.json",
    )
    args = parser.parse_args()
    docx_path = Path(args.docx)
    report_path = Path(args.report)
    doc = Document(docx_path)
    warnings: list[str] = []
    failures: list[str] = []

    headings = [
        {"style": p.style.name, "text": p.text.strip()}
        for p in doc.paragraphs
        if p.style is not None and p.style.name.startswith("Heading")
    ]
    if not headings:
        warnings.append("No Word heading styles found.")

    empty_tables = []
    uneven_tables = []
    for index, table in enumerate(doc.tables):
        if not table.rows or not table.columns:
            empty_tables.append(index)
            continue
        row_widths = [len(row.cells) for row in table.rows]
        if len(set(row_widths)) > 1:
            uneven_tables.append({"table": index, "row_cell_counts": row_widths})

    table_geometry = []
    for index, table in enumerate(doc.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find("w:tblW", NS)
        tbl_ind = tbl_pr.find("w:tblInd", NS)
        grid_cols = table._tbl.tblGrid.findall("w:gridCol", NS)
        grid_widths = [
            int(node.get(f"{{{NS['w']}}}w", "0"))
            for node in grid_cols
        ]
        row_widths = []
        for row in table.rows:
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.tcPr.find("w:tcW", NS)
                widths.append(
                    int(tc_w.get(f"{{{NS['w']}}}w", "0"))
                    if tc_w is not None
                    else 0
                )
            row_widths.append(widths)
        record = {
            "table": index,
            "tblW": (
                int(tbl_w.get(f"{{{NS['w']}}}w", "0"))
                if tbl_w is not None
                else 0
            ),
            "tblInd": (
                int(tbl_ind.get(f"{{{NS['w']}}}w", "0"))
                if tbl_ind is not None
                else 0
            ),
            "grid": grid_widths,
            "row_widths": row_widths,
        }
        table_geometry.append(record)
        if (
            record["tblW"] != 9360
            or record["tblInd"] != 120
            or sum(grid_widths) != 9360
            or any(widths != grid_widths for widths in row_widths)
        ):
            failures.append(f"Table {index} has inconsistent DXA geometry.")

    with zipfile.ZipFile(docx_path) as archive:
        names = set(archive.namelist())
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        rels_xml = etree.fromstring(archive.read("word/_rels/document.xml.rels"))

        image_members = sorted(name for name in names if name.startswith("word/media/"))
        image_rels = [
            rel.get("Target")
            for rel in rels_xml
            if rel.get("Type", "").endswith("/image")
        ]
        missing_image_targets = [
            target
            for target in image_rels
            if f"word/{target.lstrip('/')}" not in names
        ]

        page_breaks = document_xml.xpath(
            ".//w:br[@w:type='page']", namespaces=NS
        )
        section_properties = document_xml.xpath(".//w:sectPr", namespaces=NS)
        field_codes = []
        for member in sorted(
            name
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        ):
            xml = etree.fromstring(archive.read(member))
            field_codes.extend(
                "".join(node.itertext()).strip()
                for node in xml.xpath(".//w:instrText", namespaces=NS)
            )
        body_text = " ".join(document_xml.xpath(".//w:t/text()", namespaces=NS))

    if missing_image_targets:
        warnings.append("One or more embedded image relationships are broken.")
    if not any("PAGE" in code.upper() for code in field_codes):
        warnings.append("No PAGE field found.")
    if len(body_text.split()) < 2500:
        warnings.append("Document is unexpectedly short for the intended theory paper.")
    if "\ufffd" in body_text:
        failures.append("Replacement character U+FFFD found.")
    required_text = [
        "21.4. KG²RAG-style baseline при равном бюджете",
        "0.743",
        "0.494",
        "0.202",
    ]
    if "v6" in docx_path.stem.lower():
        required_text.extend(
            [
                "Версия 6.0",
                "27. Перенос на ранее не виденные official-dev запросы",
                "29. Официальный код HippoRAG на released MuSiQue",
                "0.366",
                "99.8%",
            ]
        )
    else:
        required_text.append("Версия 5.0")
    for required in required_text:
        if required not in body_text:
            failures.append(f"Required final text is missing: {required}")
    for obsolete in (
        "Не реализован опубликованный graph/iterative system baseline",
        "Published baseline и полный reader run завершены.",
        "Версия 3.0",
    ):
        if obsolete in body_text:
            failures.append(f"Obsolete text remains: {obsolete}")

    result = {
        "file": str(docx_path.resolve()),
        "size_bytes": docx_path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "headings": headings,
        "tables": len(doc.tables),
        "empty_tables": empty_tables,
        "uneven_tables": uneven_tables,
        "inline_shapes": len(doc.inline_shapes),
        "embedded_images": image_members,
        "missing_image_targets": missing_image_targets,
        "sections": len(doc.sections),
        "section_properties": len(section_properties),
        "explicit_page_breaks": len(page_breaks),
        "field_codes": field_codes,
        "word_count_approx": len(body_text.split()),
        "table_geometry": table_geometry,
        "warnings": warnings,
        "failures": failures,
        "structural_status": (
            "PASS"
            if not missing_image_targets and not empty_tables and not failures
            else "FAIL"
        ),
        "visual_render_status": "UNAVAILABLE_NO_LIBREOFFICE_OR_WORD",
    }
    report_path.write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(result, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
