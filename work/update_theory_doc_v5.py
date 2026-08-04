from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "Graph_Rescue_RAG_Theory_and_Results_v5.docx"
LEGACY_SOURCE = ROOT / "outputs" / "Graph_Rescue_RAG_Theory_and_Results_v2.docx"
SOURCE = LEGACY_SOURCE
FIGURE = ROOT / "outputs" / "final_v1" / "analysis" / "full_evidence_by_dataset.png"
FULL_READER = ROOT / "outputs" / "final_v1" / "analysis" / "reader_full_summary.json"
MULTISEED_ROBUSTNESS = (
    ROOT
    / "outputs"
    / "final_v1"
    / "analysis"
    / "robustness_multiseed_metrics.csv"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66707A"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF5D9"
GOLD = "806000"
PALE_RED = "FCE8E6"
RED = "9B1C1C"
WHITE = "FFFFFF"
BORDER = "C9D3DD"


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    for name in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{name}"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def replace_paragraph_text(paragraph, text: str, *, size=10, color=MUTED) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)


def is_page_break_paragraph(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    if "".join(element.itertext()).strip():
        return False
    return any(
        node.get(qn("w:type")) == "page" for node in element.iter(qn("w:br"))
    )


def remove_page_break_before_marker(doc: Document, marker: str) -> None:
    target = next((p for p in doc.paragraphs if p.text.startswith(marker)), None)
    if target is None:
        raise RuntimeError(f"Marker not found: {marker}")
    body = doc._body._element
    children = list(body)
    start = children.index(target._p)
    if start and is_page_break_paragraph(children[start - 1]):
        body.remove(children[start - 1])


def remove_from_marker(doc: Document, marker: str) -> None:
    target = next((p for p in doc.paragraphs if p.text.startswith(marker)), None)
    if target is None:
        raise RuntimeError(f"Marker not found: {marker}")
    body = doc._body._element
    children = list(body)
    start = children.index(target._p)
    if start and is_page_break_paragraph(children[start - 1]):
        start -= 1
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    widths_dxa = [int(round(width * 1440)) for width in widths]
    total = sum(widths_dxa)
    if total != 9360:
        raise ValueError(f"Table widths must sum to 6.5 in / 9360 DXA, got {total}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=8.3, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                shade_cell(cells[index], LIGHT_GRAY)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index > 0 and len(value) < 28
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(value)
            set_run_font(run, size=8.3)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_paragraph(doc, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        run = paragraph.add_run(bold_lead)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.widow_control = True
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.widow_control = True
        run = paragraph.add_run(item)
        set_run_font(run)


def add_callout(doc, title: str, text: str, *, kind: str = "info") -> None:
    fill, accent = {
        "info": (LIGHT_BLUE, BLUE),
        "note": (LIGHT_GRAY, DARK_BLUE),
        "warning": (PALE_GOLD, GOLD),
        "risk": (PALE_RED, RED),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, color=accent, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.3)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_figure(doc) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(FIGURE), width=Inches(6.2))
    inline._inline.docPr.set(
        "descr",
        "Столбчатая диаграмма full-evidence rate для hybrid, gated MRV и oracle на трёх датасетах.",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = False
    run = caption.add_run(
        "Рисунок 3. Full supporting-evidence retrieval: primary embedding, seed 101."
    )
    set_run_font(run, size=9, color=MUTED, italic=True)


doc = Document(SOURCE)

# The standalone references page in the legacy document leaves most of the
# preceding glossary continuation page empty in LibreOffice/Word pagination.
# Let the references follow the glossary naturally; the following Part II
# boundary remains an explicit page break.
remove_page_break_before_marker(doc, "18. Ключевые источники")

for paragraph in doc.paragraphs:
    if paragraph.text.startswith("Версия 2.0"):
        replace_paragraph_text(
            paragraph,
            "Версия 5.0  |  30 июля 2026  |  полный reader, актуальная литература и Git-релиз",
        )
        break

remove_from_marker(
    doc,
    "Часть II. Реализованный эксперимент, результаты и публикационный план",
)

doc.add_page_break()
add_heading(
    doc,
    "Часть II. Финальный экспериментальный snapshot",
    level=1,
)
add_paragraph(
    doc,
    "Эта часть заменяет устаревший пилотный отчёт версии 2.0. Ниже зафиксированы "
    "результаты чистого протокола: по 1 000 train- и 1 000 eval-вопросов на каждом "
    "датасете, три seed обучения, второй embedding-backbone, официальный reader-"
    "диагностический прогон, controlled graph corruption и KG²RAG-style "
    "equal-budget baseline.",
)
add_callout(
    doc,
    "Главный итог",
    "Selective graph rescue устойчиво улучшил полноту supporting evidence относительно "
    "hybrid baseline: +17.0 п.п. на HotpotQA, +20.5 п.п. на 2WikiMultiHopQA и "
    "+9.6 п.п. на MuSiQue. Все три paired bootstrap intervals полностью выше нуля. "
    "Это сильный результат внутри замороженного pooled-corpus протокола, но не SOTA-claim.",
    kind="info",
)

add_heading(doc, "19. Что именно реализовано", level=1)
add_heading(doc, "19.1. Flat retrieval", level=2)
add_bullets(
    doc,
    [
        "BM25 и dense retrieval; Qwen3 Embedding 0.6B является primary representation.",
        "BGE-M3 используется как независимая representation-sensitivity проверка.",
        "Reciprocal Rank Fusion и лёгкий feature reranker формируют seeds.",
        "Для всех сравниваемых политик зафиксированы одинаковые initial rankings, "
        "evidence budget, seed_k, final_k и action budget.",
        "Сохраняются per-query rankings, contexts, policy traces и checkpoints.",
    ],
)
add_heading(doc, "19.2. Source-grounded graph", level=2)
add_bullets(
    doc,
    [
        "Узлы соответствуют passages и нормализованным сущностям/заголовкам.",
        "Рёбра создаются из предоставленных entity annotations, заголовков и их упоминаний.",
        "Gold supporting facts не используются для graph construction.",
        "Каждый path содержит provenance, hop count, confidence и hubness-признаки.",
        "Максимум: 2 hops и 2 actions для HotpotQA/2Wiki; 3 hops и 3 actions для MuSiQue.",
    ],
)
add_heading(doc, "19.3. MRV и calibrated gate", level=2)
add_paragraph(
    doc,
    "Marginal Rescue Value оценивает не абсолютную похожесть passage на вопрос, а его "
    "условную добавочную ценность относительно уже выбранного evidence. Utility учитывает "
    "вероятность восстановления support, завершения evidence chain, reader gain, "
    "token/hop cost и риск noisy expansion.",
)
add_paragraph(
    doc,
    "Graph Rescuability Gate запускается до первого действия и повторно после изменения "
    "состояния. Threshold выбирается на calibration split при ограничении на recall "
    "rescuable cases. Это позволяет экономить действия, не превращая stopping policy в "
    "обычный classifier по accuracy.",
)
add_heading(doc, "19.4. Counterfactual reader supervision", level=2)
add_paragraph(
    doc,
    "Qwen3-8B генерировал counterfactual labels на фиксированной train-подвыборке: "
    "одинаковый вопрос решался по исходному evidence и после добавления одного graph "
    "candidate. Всего получено 158 строк для HotpotQA, 160 для 2Wiki и 154 для MuSiQue, "
    "по 60 уникальных вопросов на датасет. Сигнал reader gain является дополнительным "
    "признаком MRV, а не источником eval labels.",
)
add_heading(doc, "19.5. Граница новизны после работ 2026 года", level=2)
add_paragraph(
    doc,
    "После актуализации литературы нельзя заявлять, что adaptive graph retrieval или "
    "conditional traversal предложены впервые. A2RAG использует evidence-sufficiency "
    "controller и progressive cost-aware graph retrieval; CatRAG динамически меняет "
    "веса переходов и борется с hub drift; HHS-RAG принимает решения на уровне "
    "subgraph; PruneRAG управляет расширением query-decomposition tree по confidence.",
)
add_callout(
    doc,
    "Защищаемая новизна",
    "Работа изучает более узкий controlled setting: граф является repair-операцией после "
    "замороженного hybrid retriever; кандидат получает conditional marginal rescue value "
    "относительно уже выбранного evidence; preflight и continuation gates калибруются "
    "раздельно; итог проверяется при равном evidence/action budget, paired uncertainty и "
    "явном corruption-анализе.",
    kind="warning",
)
add_paragraph(
    doc,
    "На дату проверки official implementations A2RAG и HHS-RAG не найдены, а CatRAG "
    "repository прямо сообщает, что основная логика будет опубликована позднее. Поэтому "
    "они включены в related work, но не заменены неточной самодельной реализацией. "
    "Перед submission code availability нужно проверить повторно.",
)

add_heading(doc, "20. Данные, графы и воспроизводимость", level=1)
add_table(
    doc,
    ["Датасет", "Train", "Eval", "Passages", "Рёбра", "Hops/actions"],
    [
        ["HotpotQA", "1 000", "1 000", "19 189", "144 586", "2 / 2"],
        ["2Wiki", "1 000", "1 000", "11 347", "74 676", "2 / 2"],
        ["MuSiQue", "1 000", "1 000", "23 630", "180 645", "3 / 3"],
    ],
    [1.25, 0.7, 0.7, 1.05, 1.05, 1.75],
)
add_paragraph(
    doc,
    "Train и eval не пересекаются. Для каждого датасета сохранены frozen_config, "
    "protocol_manifest, input hashes, code-tree hash и leakage audit. Все три audit "
    "завершились без medium/high findings.",
)
add_table(
    doc,
    ["Датасет", "Protocol ID"],
    [
        ["HotpotQA", "02fea6d4337bd7105f38aba29ae0b13c8830d1a0031cb0c23410b3b90d70c12f"],
        ["2Wiki", "914a116c9fbf473b4d806a3525e46b8805f9cf56cb5c7f7f9174dca6ea345a32"],
        ["MuSiQue", "06f341a90ccf54c80cb8c039cee720d380dafbcf1386a2f243f7f549b6bf68e7"],
    ],
    [1.25, 5.25],
)

add_heading(doc, "21. Основные retrieval-результаты", level=1)
add_table(
    doc,
    ["Датасет", "Hybrid FE", "Gated FE", "Разность", "95% CI", "Oracle FE"],
    [
        ["HotpotQA", "0.640", "0.810", "+0.170", "[0.143; 0.197]", "0.932"],
        ["2Wiki", "0.367", "0.572", "+0.205", "[0.178; 0.231]", "0.753"],
        ["MuSiQue", "0.156", "0.252", "+0.096", "[0.074; 0.118]", "0.438"],
    ],
    [1.25, 0.85, 0.85, 0.85, 1.65, 1.05],
)
add_paragraph(
    doc,
    "FE означает долю вопросов, для которых в финальном контексте присутствует вся "
    "размеченная supporting chain. CI получены paired bootstrap по вопросам, поэтому "
    "учитывают прямое попарное сравнение двух политик на одних и тех же запросах.",
)
add_figure(doc)

add_heading(doc, "21.1. Устойчивость к seed", level=2)
add_table(
    doc,
    ["Датасет", "Hybrid FE", "Gated FE, mean ± SD", "Actions, mean ± SD"],
    [
        ["HotpotQA", "0.640", "0.806 ± 0.003", "1.652 ± 0.086"],
        ["2Wiki", "0.367", "0.573 ± 0.001", "1.565 ± 0.079"],
        ["MuSiQue", "0.156", "0.255 ± 0.003", "2.616 ± 0.087"],
    ],
    [1.3, 1.0, 2.15, 2.05],
)
add_paragraph(
    doc,
    "Низкое стандартное отклонение по seed 101/202/303 показывает, что основной "
    "retrieval-эффект не является случайностью одного обучения.",
)

add_heading(doc, "21.2. Смена embedding-backbone", level=2)
add_table(
    doc,
    ["Датасет", "Qwen H/G", "Qwen Δ", "BGE H/G", "BGE Δ"],
    [
        ["HotpotQA", "0.640 / 0.810", "+0.170", "0.669 / 0.822", "+0.153"],
        ["2Wiki", "0.367 / 0.572", "+0.205", "0.393 / 0.601", "+0.208"],
        ["MuSiQue", "0.156 / 0.252", "+0.096", "0.164 / 0.265", "+0.101"],
    ],
    [1.25, 1.55, 1.0, 1.55, 1.15],
)
add_paragraph(
    doc,
    "BGE-M3 проверялся на seed 101. Направление и величина эффекта близки к primary "
    "backbone; это подтверждает representation robustness, но не заменяет полноценные "
    "три seed для BGE.",
)

add_heading(doc, "21.3. Wins, losses и ties", level=2)
add_table(
    doc,
    ["Датасет", "Wins", "Losses", "Ties"],
    [
        ["HotpotQA", "199", "33", "768"],
        ["2Wiki", "328", "19", "653"],
        ["MuSiQue", "164", "81", "755"],
        ["Всего", "691", "133", "2 176"],
    ],
    [2.0, 1.5, 1.5, 1.5],
)
add_paragraph(
    doc,
    "Wins значительно больше losses, однако 133 ухудшения подтверждают, что graph "
    "expansion не является безусловно безопасным. Именно поэтому selector и gate — "
    "центральные компоненты, а не декоративная оптимизация.",
)

add_heading(doc, "21.4. KG²RAG-style baseline при равном бюджете", level=2)
add_paragraph(
    doc,
    "Реализована независимая адаптация опубликованного паттерна KG²RAG: semantic seeds "
    "расширяются через граф, кандидаты получают query relevance, propagated seed score "
    "и bonus за поддержку несколькими seeds, после чего контекст организуется в "
    "seed-центричные группы. Это не точное воспроизведение оригинальной системы: "
    "исходные triplet KG, relation extraction и FlagReranker заменены нашим "
    "passage/entity graph и общим semantic scorer.",
)
add_table(
    doc,
    [
        "Датасет",
        "Hybrid FE",
        "KG²-style FE",
        "Gated FE",
        "KG²−Hybrid",
        "Gated−KG²",
    ],
    [
        ["HotpotQA", "0.640", "0.743", "0.810", "+0.103", "+0.067"],
        ["2Wiki", "0.367", "0.494", "0.572", "+0.127", "+0.078"],
        ["MuSiQue", "0.156", "0.202", "0.252", "+0.046", "+0.050"],
    ],
    [1.25, 1.0, 1.15, 1.0, 1.05, 1.05],
)
add_paragraph(
    doc,
    "Для KG²-style против hybrid paired 95% CI равны [0.080; 0.127], "
    "[0.105; 0.150] и [0.027; 0.065]. Для KG²-style минус gated MRV интервалы "
    "равны [−0.089; −0.045], [−0.102; −0.053] и [−0.067; −0.034]. "
    "Все интервалы исключают ноль. В 3 000 запросах не найдено ни одного "
    "нарушения final_k, token budget или action budget и ни одного расхождения "
    "initial seeds с замороженным run.",
)
add_callout(
    doc,
    "Что теперь можно утверждать",
    "Общий паттерн seed → graph expansion действительно сильнее flat hybrid retrieval. "
    "Однако gated MRV стабильно лучше KG²RAG-style control при тех же бюджетах. "
    "Следовательно, преимущество Graph Rescue нельзя объяснить только безусловным "
    "добавлением графовых соседей и организацией контекста; дополнительную ценность "
    "дают conditional gate и marginal-value selector.",
    kind="info",
)

add_heading(doc, "22. Gate и эффективность", level=1)
add_table(
    doc,
    ["Датасет", "AUROC", "ECE", "Gate positive", "Gated actions"],
    [
        ["HotpotQA", "0.736", "0.049", "0.914", "1.57"],
        ["2Wiki", "0.813", "0.055", "0.832", "1.52"],
        ["MuSiQue", "0.651", "0.068", "0.991", "2.69"],
    ],
    [1.45, 1.0, 1.0, 1.55, 1.5],
)
add_paragraph(
    doc,
    "Относительно MRV-always gate уменьшает среднее число graph actions примерно на "
    "20% для HotpotQA, 23% для 2Wiki и 10% для MuSiQue. Потеря full evidence мала. "
    "MuSiQue остаётся слабым местом: gate пропускает 99.1% запросов и почти не выполняет "
    "функцию preflight routing.",
)
add_callout(
    doc,
    "Что гипотеза о скорости пока не доказала",
    "Снижение graph actions относительно always-expand proxy ещё не доказывает, что "
    "система быстрее полного GraphRAG. Не измерены одинаковым способом offline graph "
    "construction, indexing, full retrieval и generation, и нет прямого GraphRAG "
    "baseline на том же hardware.",
    kind="warning",
)

add_heading(doc, "23. Qwen3-8B reader: полный downstream-прогон", level=1)
reader_summary = json.loads(FULL_READER.read_text(encoding="utf-8"))
reader_names = {"hotpot": "HotpotQA", "2wiki": "2Wiki", "musique": "MuSiQue"}
reader_rows = []
for reader_result in reader_summary["results"]:
    hybrid = reader_result["hybrid"]
    gated = reader_result["mrv_gated"]
    delta = reader_result["deltas"]["answer_f1"]
    reader_rows.append(
        [
            reader_names[reader_result["dataset"]],
            str(reader_result["queries"]),
            f'{hybrid["answer_f1"]:.3f} / {gated["answer_f1"]:.3f}',
            f'{hybrid["support_f1"]:.3f} / {gated["support_f1"]:.3f}',
            (
                f'{delta["difference"]:+.3f} '
                f'[{delta["ci95_low"]:+.3f}; {delta["ci95_high"]:+.3f}]'
            ),
        ]
    )
add_paragraph(
    doc,
    "Один и тот же локальный Qwen3-8B reader, prompt, decoding и token budget применены "
    "к контекстам hybrid и gated MRV. Для scoring использованы официальные evaluator "
    "implementations; интервалы рассчитаны попарно по одним и тем же вопросам из "
    "официальных per-query определений метрик, а p-values скорректированы методом Холма "
    "внутри каждого датасета.",
)
add_table(
    doc,
    ["Датасет", "N", "Answer F1 H/G", "Support F1 H/G", "Δ Answer F1 [95% CI]"],
    reader_rows,
    [1.1, 0.55, 1.35, 1.4, 2.1],
)
if reader_summary["complete"]:
    add_paragraph(
        doc,
        "Полный reader-анализ завершён на 1 000 eval-вопросах каждого датасета. "
        "Answer F1 значимо вырос на HotpotQA и 2Wiki; положительный эффект MuSiQue "
        "не сохраняет значимость после поправки Холма, а Support F1 там практически "
        "не изменился. Это позволяет делать ограниченный внутрипроектный claim о "
        "downstream-переносе на двух наборах, но не превращает pooled-corpus setting "
        "в официальный leaderboard benchmark.",
    )
else:
    add_callout(
        doc,
        "Промежуточное состояние",
        "На момент сборки документа полный reader-прогон ещё не завершён для: "
        + ", ".join(reader_names[name] for name in reader_summary["datasets_missing"])
        + ". В таблицу включены только полностью завершённые датасеты.",
        kind="warning",
    )
add_callout(
    doc,
    "Почему 8B-модель допустима",
    "Локальный 8B reader подходит для controlled resource-bounded comparison: политикам "
    "выдана одна и та же модель, prompt, decoding и token budget. Можно сравнивать "
    "внутренние изменения. Нельзя сопоставлять абсолютный Answer F1 с работами на "
    "GPT-4/70B и объяснять разницу retrieval-методом.",
    kind="note",
)

add_heading(doc, "24. Устойчивость к повреждению графа", level=1)
robustness_headers = ["Датасет", "Clean", "D10", "D25", "D50", "F10", "F25", "F50", "Mix"]
robustness_widths = [1.1, 0.675, 0.675, 0.675, 0.675, 0.675, 0.675, 0.675, 0.675]
if MULTISEED_ROBUSTNESS.exists():
    with MULTISEED_ROBUSTNESS.open(encoding="utf-8", newline="") as handle:
        robustness_records = [
            row for row in csv.DictReader(handle) if row["policy"] == "mrv_gated"
        ]
    robustness_index = {
        (row["dataset"], row["condition"]): row for row in robustness_records
    }
    condition_order = [
        "clean",
        "dropout_10",
        "dropout_25",
        "dropout_50",
        "false_edges_10",
        "false_edges_25",
        "false_edges_50",
        "mixed_25_25",
    ]
    robustness_rows = []
    for dataset, label in (("hotpot", "Hotpot"), ("2wiki", "2Wiki"), ("musique", "MuSiQue")):
        cells = [label]
        for condition in condition_order:
            record = robustness_index[(dataset, condition)]
            mean_value = float(record["full_evidence_mean"])
            std_value = float(record["full_evidence_std"])
            if int(record["seeds"]) <= 1:
                cells.append(f"{mean_value:.3f}")
            else:
                cells.append(f"{mean_value:.3f}±{std_value:.3f}")
        robustness_rows.append(cells)
    robustness_note = (
        "D — edge dropout; F — добавление ложных рёбер; Mix — D25/F25. "
        "В ячейках показаны mean±SD по пяти детерминированным corruption seeds "
        "(для clean-графа seed не влияет). Stress-test jobs выполнялись параллельно, "
        "поэтому их latency не используется для сравнительных выводов."
    )
    robustness_caution = (
        "Multi-seed агрегация завершена; следующий эксперимент должен проверить "
        "обучаемую edge-confidence/denoising стратегию."
    )
else:
    robustness_rows = [
        ["Hotpot", "0.810", "0.770", "0.757", "0.686", "0.668", "0.580", "0.537", "0.574"],
        ["2Wiki", "0.572", "0.544", "0.522", "0.461", "0.450", "0.348", "0.316", "0.342"],
        ["MuSiQue", "0.252", "0.228", "0.251", "0.189", "0.198", "0.148", "0.129", "0.145"],
    ]
    robustness_note = (
        "D — edge dropout; F — добавление ложных рёбер; Mix — D25/F25. "
        "Это предварительный single-seed stress test; multi-seed оценка ещё не завершена."
    )
    robustness_caution = (
        "Вывод о форме dose-response следует считать окончательным только после "
        "multi-seed агрегации."
    )
add_table(
    doc,
    robustness_headers,
    robustness_rows,
    robustness_widths,
)
add_paragraph(
    doc,
    robustness_note + " Модели обучались только на clean graph. False edges во всех "
    "случаях вредят сильнее dropout и дают монотонную dose-response деградацию.",
)
add_callout(
    doc,
    "Следующий методический приоритет",
    "Нужен precision-aware edge-confidence/denoising слой. Результаты не поддерживают "
    "стратегию «сначала сделать граф максимально полным»: ложная связь опаснее "
    "отсутствующей. " + robustness_caution,
    kind="info",
)

add_heading(doc, "25. Адекватны ли данные и метрики", level=1)
add_heading(doc, "25.1. Для чего текущий протокол адекватен", level=2)
add_bullets(
    doc,
    [
        "Три признанных multi-hop QA benchmark имеют supporting-fact/passage разметку.",
        "1 000 eval-вопросов на датасет дают информативные paired confidence intervals.",
        "Методы сравниваются на одинаковых запросах, initial rankings и бюджетах.",
        "Train/calibration/eval разделены; протоколы и hashes зафиксированы.",
        "Есть три seed, второй embedding-backbone, oracle и controlled corruption.",
        "Full-evidence rate и support recall напрямую измеряют заявленный retrieval bottleneck.",
    ],
)
add_heading(doc, "25.2. Для чего протокол не адекватен", level=2)
add_bullets(
    doc,
    [
        "Это pooled corpus из passages выбранных вопросов, а не официальный leaderboard setting.",
        "HotpotQA не является ни standard distractor submission, ни fullwiki submission.",
        "Абсолютные проценты нельзя напрямую сравнивать с таблицами других работ.",
        (
            "Полный reader-прогон уменьшает статистическую неопределённость downstream-оценки, "
            "но слабая 8B-модель и pooled corpus не позволяют заявлять leaderboard/SOTA answer quality."
            if reader_summary["complete"]
            else "Reader-прогон ещё не завершён на всех трёх датасетах."
        ),
        "KG²RAG-style control не является точным запуском оригинального triplet-KG pipeline.",
        "Не измерен полный offline/online GraphRAG lifecycle.",
    ],
)
add_paragraph(
    doc,
    "Следовательно, данных уже достаточно для controlled retrieval paper, но не для "
    "leaderboard/SOTA paper. Наиболее честная формулировка — pooled-corpus multi-hop "
    "retrieval study derived from public benchmarks.",
)

add_heading(doc, "26. Разрешённые и запрещённые claims", level=1)
add_table(
    doc,
    ["Допустимо", "Недопустимо"],
    [
        [
            "Gated graph rescue повысил full evidence внутри замороженного протокола.",
            "Метод достиг SOTA на HotpotQA/2Wiki/MuSiQue.",
        ],
        [
            "Эффект устойчив к seed и смене Qwen/BGE embeddings.",
            "Результат не зависит от reader или embedding-модели вообще.",
        ],
        [
            (
                "8B reader показал положительный paired downstream signal на 1 000 вопросах "
                "каждого датасета."
                if reader_summary["complete"]
                else "8B reader дал положительный сигнал на завершённых частях полного прогона."
            ),
            "Метод достиг SOTA answer quality в официальном leaderboard setting.",
        ],
        [
            "False edges опаснее missing edges в проведённом corruption test.",
            "Любой knowledge graph обязан вести себя так же.",
        ],
        [
            "Gate уменьшил actions относительно always-expand proxy.",
            "Система быстрее полного GraphRAG без прямого сравнения.",
        ],
        [
            "Gated MRV выше независимой KG²RAG-style adaptation при равном бюджете.",
            "Наш метод превосходит оригинальную опубликованную реализацию KG²RAG.",
        ],
    ],
    [3.25, 3.25],
)

doc.add_page_break()
add_heading(doc, "Часть III. Что делать дальше и как публиковать", level=1)
add_heading(doc, "27. Следующие обязательные научные шаги", level=1)
add_numbered(
    doc,
    [
        (
            "Полный Qwen3-8B reader-прогон на 3 × 1 000 вопросах завершён; перед релизом "
            "нужно проверить агрегаты и сверить выборочные raw generations."
            if reader_summary["complete"]
            else "Завершить Qwen3-8B reader на 1 000 eval-вопросах каждого датасета."
        ),
        (
            "Multi-seed graph corruption завершён; проверить агрегацию и отразить "
            "mean±SD в статье."
            if MULTISEED_ROBUSTNESS.exists()
            else "Повторить graph corruption с пятью seeds и агрегировать mean±SD."
        ),
        "Провести ручной error analysis не менее 100 traces, стратифицированных по "
        "win/loss/tie, anchor failure, graph noise и gate false negative.",
        "Измерить offline graph/index construction и p50/p95 end-to-end latency в "
        "одинаковом hardware/software окружении.",
    ],
)
add_paragraph(
    doc,
    "Published-pattern baseline закрыт. Главные оставшиеся риски — отсутствие независимого "
    "человеческого trace audit и полного end-to-end cost comparison с GraphRAG на том же "
    "оборудовании. Новые механизмы до закрытия этих пунктов скорее размоют статью, чем усилят её.",
)

add_heading(doc, "28. Рекомендуемый маршрут публикации", level=1)
add_heading(doc, "28.1. Journal of Intelligent Information Systems", level=2)
add_paragraph(
    doc,
    "Наиболее точное тематическое совпадение: intelligent information retrieval на стыке "
    "AI и database systems, включая graph-based retrieval. Журнал hybrid: при обычном "
    "subscription-маршруте обязательного APC нет; open-access вариант платный. Это более "
    "строгий вариант, особенно после появления близких adaptive graph retrieval работ 2026 года, "
    "поэтому novelty должна быть сформулирована как local repair after frozen hybrid retrieval, "
    "paired equal-budget evaluation и анализ false-vs-missing edges.",
)
add_paragraph(
    doc,
    "Overview: https://link.springer.com/journal/10844\n"
    "Publishing options: https://link.springer.com/journal/10844/how-to-publish-with-us",
)
add_heading(doc, "28.2. SN Computer Science", level=2)
add_paragraph(
    doc,
    "Более широкий резервный journal route без поездки: scope включает AI, NLP и "
    "information retrieval. Рассматривать его следует только после решения JIIS и "
    "повторной проверки актуального квартиля, индексации и publishing model.",
)
add_paragraph(
    doc,
    "Overview: https://link.springer.com/journal/42979\n"
    "Publishing options: https://link.springer.com/journal/42979/how-to-publish-with-us",
)
add_heading(doc, "28.3. Запасные журнальные варианты", level=2)
add_paragraph(
    doc,
    "Discover Artificial Intelligence тематически подходит для AI methodology и knowledge "
    "reasoning, а Discover Computing — для более широкого retrieval/system framing. Оба "
    "варианта следует рассматривать после проверки актуального APC, waiver и indexing.",
)
add_paragraph(
    doc,
    "Discover AI: https://link.springer.com/journal/44163/aims-and-scope\n"
    "Discover Computing: https://link.springer.com/journal/10791/aims-and-scope",
)
add_heading(doc, "28.4. GitHub, Zenodo и arXiv", level=2)
add_paragraph(
    doc,
    "Практический порядок: публичный GitHub release с tag v0.1.0 → архивирование release "
    "в Zenodo и получение DOI → добавление DOI в CITATION.cff и manuscript → arXiv preprint "
    "после человеческой проверки → journal submission. arXiv moderation не является peer "
    "review; для первой категории может понадобиться endorsement.",
)
add_paragraph(
    doc,
    "Zenodo GitHub integration: https://help.zenodo.org/docs/github/\n"
    "Moderation: https://info.arxiv.org/help/moderation/index.html\n"
    "Endorsement: https://info.arxiv.org/help/endorsement.html",
)
add_heading(doc, "28.5. Русскоязычный маршрут", level=2)
add_paragraph(
    doc,
    "Русский текст сохраняется как подробная рабочая теория и материал для проверки "
    "аргументации. До решения JIIS его не следует публиковать отдельной полной статьёй, "
    "чтобы не создавать риск duplicate publication и не лишать английскую рукопись "
    "статуса оригинальной подачи.",
)
add_callout(
    doc,
    "Нельзя просто перевести опубликованную статью",
    "Полная русская публикация, затем её перевод в иностранном журнале без согласия "
    "обеих редакций и явного раскрытия является duplicate publication. Для этого проекта "
    "первичной работой выбрана английская рукопись JIIS; русский документ остаётся "
    "неархивным рабочим объяснением. Отдельная последующая статья должна содержать новый "
    "вопрос, новые эксперименты и явную ссылку на первую работу.",
    kind="warning",
)
add_paragraph(
    doc,
    "Наиболее естественная самостоятельная тема международного продолжения — "
    "precision-aware graph rescue under uncertain edges: новый edge-confidence/denoising "
    "механизм, open-domain или прикладной corpus, одинаково измеренный полный GraphRAG "
    "lifecycle cost и новые public baselines. Это будет follow-up study, а не перевод.",
)
add_paragraph(
    doc,
    "COPE duplicate publication: https://doi.org/10.24318/y9lyqPiR",
)

add_heading(doc, "29. Итоговая формулировка проекта", level=1)
add_callout(
    doc,
    "Наиболее защищаемый научный вывод",
    "В pooled-corpus multi-hop retrieval selective local graph expansion, управляемый "
    "conditional marginal-value selector и calibrated gate, повышает полноту supporting "
    "evidence относительно сильного hybrid baseline на трёх датасетах. Эффект устойчив "
    "к seed и embedding representation, но чувствителен прежде всего к ложным рёбрам. "
    "Независимый KG²RAG-style control подтверждает пользу общего graph-expansion pattern, "
    "но gated MRV остаётся выше при равном бюджете. Локальный 8B reader подтверждает "
    + (
        "paired downstream-перенос на 3 × 1 000 вопросах."
        if reader_summary["complete"]
        else "предварительный downstream-перенос; полный прогон ещё не завершён."
    ),
    kind="info",
)
add_paragraph(
    doc,
    "Работа уже имеет самостоятельное содержание: не просто комбинацию RAG и графа, а "
    "исследование того, когда локальный графовый обход полезен, как оценивать условную "
    "ценность кандидата и как ложные связи меняют результат. Следующий этап — не добавлять "
    "ещё один механизм, а завершить ручной trace audit, end-to-end cost benchmark и "
    "человеческую ревизию текста до публикационного уровня.",
)

doc.core_properties.title = "Graph Rescue RAG: теория, реализация и итоговые результаты"
doc.core_properties.subject = "Selective graph expansion for multi-hop retrieval"
doc.core_properties.comments = (
    "Version 5.0. Adds full-reader evaluation, 2026 novelty boundary, release guidance, "
    "and multi-seed corruption results when available."
)
doc.save(OUTPUT)
print(OUTPUT)
