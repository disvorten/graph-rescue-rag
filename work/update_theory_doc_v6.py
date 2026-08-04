from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "Graph_Rescue_RAG_Theory_and_Results_v5.docx"
OUTPUT = ROOT / "outputs" / "Graph_Rescue_RAG_Theory_and_Results_v6.docx"
GLOBAL = ROOT / "outputs" / "global_v1" / "analysis" / "analysis_summary.json"
GATE = ROOT / "outputs" / "global_v1" / "gate_transfer" / "analysis_summary.json"
OFFICIAL = (
    ROOT
    / "outputs"
    / "official_baselines"
    / "aligned_analysis"
    / "analysis_summary.json"
)
LATENCY_ROOT = ROOT / "outputs" / "latency_v1"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF5D9"
MUTED = "66707A"
WHITE = "FFFFFF"


def read_required(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"Required final result is missing: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def set_font(run, *, size=10.0, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    for name in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{name}"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def replace_paragraph(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    set_font(paragraph.add_run(text), size=10.5, color=MUTED)


def is_page_break_paragraph(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    if "".join(element.itertext()).strip():
        return False
    return any(
        node.get(qn("w:type")) == "page" for node in element.iter(qn("w:br"))
    )


def remove_from_marker(doc: Document, marker: str) -> None:
    target = next((p for p in doc.paragraphs if p.text.startswith(marker)), None)
    if target is None:
        raise RuntimeError(f"Marker not found: {marker}")
    body = doc._body._element
    children = list(body)
    start = children.index(target._p)
    if start and is_page_break_paragraph(children[start - 1]):
        start -= 1
    for child in children[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_font(
            run,
            size=15 if level == 1 else 12,
            color=DARK_BLUE,
            bold=True,
        )
    return paragraph


def paragraph(doc: Document, text: str, *, muted: bool = False) -> None:
    item = doc.add_paragraph()
    item.paragraph_format.space_after = Pt(6)
    item.paragraph_format.line_spacing = 1.08
    set_font(item.add_run(text), size=10, color=MUTED if muted else None)


def callout(doc: Document, title: str, text: str, *, warning: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), PALE_GOLD if warning else LIGHT_BLUE)
    cell._tc.get_or_add_tcPr().append(shade)
    p = cell.paragraphs[0]
    set_font(p.add_run(title + "\n"), size=10, color=DARK_BLUE, bold=True)
    set_font(p.add_run(text), size=9.5)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def shade_cell(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, *, top=60, start=70, bottom=60, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    raw = [max(width, 0.01) for width in widths]
    scale = 9360 / sum(raw)
    widths_dxa = [int(round(width * scale)) for width in raw]
    widths_dxa[-1] += 9360 - sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    caption: str,
    widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.keep_with_next = True
    set_font(cap.add_run(caption), size=9, color=MUTED, italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.keep_with_next = True
        set_font(cell.paragraphs[0].add_run(value), size=8.5, color=WHITE, bold=True)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_properties.append(table_header)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                shade_cell(cells[index], LIGHT_GRAY)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cells[index].paragraphs[0].add_run(value), size=8.5)
    if widths:
        set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def pct(value: float) -> str:
    return f"{100.0 * value:.1f}%"


def f3(value: float) -> str:
    return f"{value:.3f}"


def global_rows(global_result: dict) -> list[list[str]]:
    names = {"hotpot": "HotpotQA", "2wiki": "2Wiki", "musique": "MuSiQue"}
    rows = []
    for dataset in ("hotpot", "2wiki", "musique"):
        item = global_result["datasets"][dataset]
        comp = item["comparisons"]["mrv_gated_vs_hybrid_full_evidence"]
        rows.append(
            [
                names[dataset],
                f"{item['queries']:,}".replace(",", " "),
                f"{item['corpus_passages']:,}".replace(",", " "),
                f3(item["policies"]["hybrid"]["full_evidence"]),
                f3(item["policies"]["mrv_gated"]["full_evidence"]),
                f"{comp['difference']:+.3f}",
                f"[{comp['ci95_low']:.3f}; {comp['ci95_high']:.3f}]",
                pct(item["gate_open_rate"]),
            ]
        )
    return rows


def gate_rows(gate_result: dict) -> list[list[str]]:
    names = {"hotpot": "HotpotQA", "2wiki": "2Wiki", "musique": "MuSiQue"}
    rows = []
    for dataset in ("hotpot", "2wiki", "musique"):
        item = gate_result["datasets"][dataset]
        frozen = item["frozen_gate_on_heldout"]
        recal = item["recalibrated_gate_on_heldout"]
        frozen_policy = item["preflight_only_policy_frozen"]
        recal_policy = item["preflight_only_policy_recalibrated"]
        rows.append(
            [
                names[dataset],
                f"{item['calibration_queries']}/{item['heldout_test_queries']}",
                f"{frozen['recall']:.3f} → {recal['recall']:.3f}",
                f"{frozen['ece']:.3f} → {recal['ece']:.3f}",
                f"{frozen_policy['open_rate']:.3f} → {recal_policy['open_rate']:.3f}",
                f"{frozen_policy['graph_actions']:.2f} → {recal_policy['graph_actions']:.2f}",
                f"{frozen_policy['full_evidence']:.3f} → {recal_policy['full_evidence']:.3f}",
            ]
        )
    return rows


def official_rows(result: dict) -> list[list[str]]:
    names = {
        "StandardRAG_official_code": "StandardRAG (official code)",
        "HippoRAG_official_code": "HippoRAG (official code)",
        "GraphRescue_hybrid": "Graph Rescue: hybrid",
        "GraphRescue_gated_MRV": "Graph Rescue: gated MRV",
    }
    rows = []
    for key, item in result["systems"].items():
        rows.append(
            [
                names.get(key, key),
                f3(item["full_evidence_at_7"]),
                f3(item["support_recall_at_7"]),
                f"{item['retrieval_latency']['median_ms']:.1f}",
                f"{item['retrieval_latency']['p95_ms']:.1f}",
            ]
        )
    return rows


def official_comparison_rows(result: dict) -> list[list[str]]:
    names = (
        (
            "GraphRescue_gated_minus_HippoRAG_full_evidence_at_7",
            "Gated MRV − HippoRAG",
        ),
        (
            "GraphRescue_gated_minus_StandardRAG_full_evidence_at_7",
            "Gated MRV − StandardRAG",
        ),
        (
            "GraphRescue_gated_minus_GraphRescue_hybrid_full_evidence_at_7",
            "Gated MRV − shared hybrid",
        ),
    )
    rows = []
    for key, label in names:
        item = result["paired_full_evidence_comparisons"][key]
        rows.append(
            [
                label,
                f"{item['difference']:+.3f}",
                f"[{item['ci95_low']:.3f}; {item['ci95_high']:.3f}]",
            ]
        )
    return rows


def latency_rows() -> list[list[str]]:
    names = {"hotpot": "HotpotQA", "2wiki": "2Wiki", "musique": "MuSiQue"}
    rows = []
    for dataset in ("hotpot", "2wiki", "musique"):
        result = read_required(LATENCY_ROOT / f"{dataset}.json")
        base = result["aggregate"]["hybrid"]
        always = result["aggregate"]["mrv_always"]
        gated = result["aggregate"]["mrv_gated"]
        comparison = result["paired_latency_comparisons"][
            "mrv_gated_minus_mrv_always_online_total_ms"
        ]
        rows.append(
            [
                names[dataset],
                f"{base['online_total_latency']['median_ms']:.1f}",
                f"{always['online_total_latency']['median_ms']:.1f}",
                f"{gated['online_total_latency']['median_ms']:.1f}",
                f"{gated['online_total_latency']['p95_ms']:.1f}",
                f"{comparison['difference']:+.1f}",
                f"[{comparison['ci95_low']:.1f}; {comparison['ci95_high']:.1f}]",
                f"{always['mean_graph_actions']:.2f} → {gated['mean_graph_actions']:.2f}",
            ]
        )
    return rows


def main() -> None:
    global_result = read_required(GLOBAL)
    gate_result = read_required(GATE)
    official_result = read_required(OFFICIAL)
    doc = Document(SOURCE)
    for item in doc.paragraphs:
        if item.text.startswith("Версия 5.0"):
            replace_paragraph(
                item,
                "Версия 6.0  |  4 августа 2026  |  пакет подготовки к JIIS",
            )
            break
    remove_from_marker(doc, "Часть III. Что делать дальше и как публиковать")

    doc.add_page_break()
    heading(doc, "Часть III. Расширенная проверка и подготовка к JIIS", 1)
    paragraph(
        doc,
        "Эта часть добавляет проверку переноса на всех ранее не виденных official-dev "
        "запросах, label-efficient recalibration gate, запуск официального кода "
        "HippoRAG и отдельный последовательный benchmark времени. Основной controlled "
        "pooled-corpus эксперимент из части II сохраняется как анализ механизма.",
    )
    callout(
        doc,
        "Граница утверждений",
        "Расширенный корпус является global development/distractor pool, а не full-wiki. "
        "Официальный HippoRAG запускается с локальными Qwen-моделями и released OpenIE, "
        "поэтому это official-code local-model reproduction, а не воспроизведение чисел статьи.",
        warning=True,
    )

    heading(doc, "27. Перенос на ранее не виденные official-dev запросы", 1)
    paragraph(
        doc,
        "Модели MRV, calibrators и thresholds seed 101 были заморожены до запуска. "
        "В граф и индекс вошли полные development/distractor passage pools и frozen "
        "train sample; evaluation IDs, использованные в прежних пилотах, исключены. "
        "Gold support, answers и decompositions не использовались при построении графа.",
    )
    add_table(
        doc,
        ["Датасет", "Queries", "Passages", "Hybrid FE", "Gated FE", "Δ", "95% CI", "Gate open"],
        global_rows(global_result),
        caption="Таблица 10. Frozen-model global-development transfer; not full-wiki.",
        widths=[0.8, 0.65, 0.75, 0.7, 0.7, 0.5, 1.25, 0.75],
    )
    paragraph(
        doc,
        "Эта таблица проверяет не только качество retrieval, но и переносимость "
        "решения OPEN/CLOSED. Положительный Δ full evidence при замороженном gate "
        "сохраняется на всех трёх датасетах, хотя его величина отличается от "
        "pooled-corpus результата. Различия open rate и recall между датасетами "
        "показывают calibration shift.",
    )

    recalibration_heading = heading(doc, "28. Label-efficient recalibration gate", 1)
    recalibration_heading.paragraph_format.page_break_before = True
    paragraph(
        doc,
        "Вторичный протокол резервирует 200 target queries для калибровки и считает "
        "метрики только на оставшейся disjoint части. Oracle target реконструируется "
        "тем же графом, hops и frontier cap и численно сверяется с основным evaluator. "
        "Retrieval replay является preflight-only ablation: при OPEN выбирается уже "
        "посчитанный MRV-always trace, при CLOSED — hybrid trace.",
    )
    add_table(
        doc,
        ["Датасет", "Cal/Test", "Recall", "ECE", "Open rate", "Actions", "Full evidence"],
        gate_rows(gate_result),
        caption="Таблица 11. Frozen → target-recalibrated gate на held-out запросах.",
        widths=[0.85, 0.75, 0.95, 0.95, 1.0, 0.9, 1.1],
    )
    callout(
        doc,
        "Интерпретация trade-off",
        "Recalibration может вернуть требуемый recall и улучшить ECE, но обычно "
        "изменяет долю открытий и число graph actions. На HotpotQA и 2Wiki open rate "
        "возрастает до 0.897 и 0.935; для MuSiQue он остаётся 0.928. Поэтому gate "
        "нельзя оценивать только по AUROC или качеству ответа: необходим совместный "
        "quality–compute profile.",
    )

    heading(doc, "29. Официальный код HippoRAG на released MuSiQue", 1)
    paragraph(
        doc,
        f"Сравнение использует {official_result['queries']} уникальных общих query IDs, "
        f"{official_result['corpus_passages']} released passages и k=7. "
        "StandardRAG и HippoRAG запускаются из зафиксированного официального commit; "
        "Graph Rescue использует тот же corpus/query adapter. Recognition memory Qwen3 "
        "работает через OpenAI-compatible endpoint Ollama с reasoning_effort=none. "
        "Все процессы выполняются последовательно на одном ноутбуке.",
    )
    add_table(
        doc,
        ["Система", "FE@7", "Support recall@7", "Median, ms", "p95, ms"],
        official_rows(official_result),
        caption="Таблица 12. Выровненный official-code/local-model baseline.",
        widths=[2.05, 0.75, 1.2, 1.0, 1.0],
    )
    add_table(
        doc,
        ["Парное сравнение", "Δ FE@7", "95% CI"],
        official_comparison_rows(official_result),
        caption="Таблица 12a. Парные различия full evidence на aligned query IDs.",
        widths=[2.8, 1.0, 1.7],
    )
    paragraph(doc, official_result["comparability_note"], muted=True)
    official_gated = official_result["systems"]["GraphRescue_gated_MRV"]
    official_delta = official_result["paired_full_evidence_comparisons"][
        "GraphRescue_gated_minus_HippoRAG_full_evidence_at_7"
    ]
    paragraph(
        doc,
        f"Gated MRV превосходит HippoRAG по FE@7 на {official_delta['difference']:+.3f} "
        f"(95% CI [{official_delta['ci95_low']:.3f}; {official_delta['ci95_high']:.3f}]). "
        f"Но gate открывается для {100 * official_gated['graph_open_rate']:.1f}% запросов "
        f"и выполняет в среднем {official_gated['mean_graph_actions']:.3f} actions. "
        "Следовательно, этот перенос подтверждает качество retrieval, но не "
        "селективность gate.",
    )

    heading(doc, "30. Время и ресурсы", 1)
    paragraph(
        doc,
        "Clean latency benchmark исключает warm-up, принудительно пересчитывает query "
        "embedding, рандомизирует query/policy order и использует три повтора по 200 "
        "запросов. Answer generation не входит в online retrieval boundary. Passage "
        "cache остаётся тёплым; initialization/indexing приводятся отдельно.",
    )
    add_table(
        doc,
        ["Датасет", "Hybrid p50", "Always p50", "Gated p50", "Gated p95", "Mean Δ G−A", "95% CI", "Actions A→G"],
        latency_rows(),
        caption="Таблица 13. Последовательный online retrieval benchmark, ms.",
        widths=[0.75, 0.7, 0.75, 0.75, 0.75, 0.75, 1.0, 0.9],
    )
    paragraph(
        doc,
        "Допустимый temporal claim относится к calibrated gate относительно always-on "
        "MRV на той же реализации: paired mean уменьшается на 6.6 мс для HotpotQA, "
        "12.0 мс для 2Wiki и 6.5 мс для MuSiQue; все интервалы исключают ноль. "
        "Экономия умеренная, поскольку значительную часть времени занимает query "
        "embedding. Результат не означает, что Graph Rescue быстрее полного GraphRAG "
        "или HippoRAG lifecycle без выровненного построения графа и generation.",
    )

    heading(doc, "31. Что в итоге можно утверждать", 1)
    callout(
        doc,
        "Основной вывод",
        "Selective local graph rescue повышает вероятность собрать полную evidence chain "
        "после сильного hybrid retrieval в controlled pooled-corpus и в расширенном "
        "unseen-dev протоколе. Conditional MRV полезнее relevance-only выбора, а gate "
        "уменьшает обход относительно always-on политики. Перенос gate требует отдельной "
        "оценки и иногда небольшой target recalibration.",
    )
    paragraph(
        doc,
        "Не подтверждены SOTA, универсальное превосходство над GraphRAG-family systems, "
        "full-Wikipedia retrieval и независимость от reader. Отрицательные результаты "
        "сохраняются: MuSiQue reader gain после Holm correction незначим, false edges "
        "вреднее dropout, а часть anchor/unreachable failures не устраняется обходом.",
    )

    heading(doc, "32. Пакет подачи в Journal of Intelligent Information Systems", 1)
    paragraph(
        doc,
        "Первичная статья готовится на английском в Springer svjour3/smallcondensed. "
        "Лимит JIIS — 25 страниц вместе с таблицами, рисунками и ссылками; допустим только "
        "LaTeX. Аннотация должна занимать 150–250 слов, keywords — 4–6. В submission "
        "передаются исходники, style files, figures и собранный PDF без подкаталогов.",
    )
    paragraph(
        doc,
        "Рабочее название: Selective Local Graph Rescue after Hybrid Retrieval: "
        "Calibrated Evidence Completion for Multi-Hop Question Answering. Позиционирование "
        "делает акцент на intelligent information retrieval, uncertainty-aware decision "
        "и evidence completion, а не на общем ярлыке GraphRAG.",
    )
    paragraph(
        doc,
        "Официальные требования: https://link.springer.com/journal/10844/submission-guidelines\n"
        "Scope: https://link.springer.com/journal/10844/aims-and-scope\n"
        "Код: https://github.com/disvorten/graph-rescue-rag\n"
        "DOI: https://doi.org/10.5281/zenodo.21709269",
        muted=True,
    )

    heading(doc, "33. Финальная граница готовности", 1)
    paragraph(
        doc,
        "Техническая готовность означает: все три global runs завершены; official-code "
        "baseline и latency benchmark имеют compact summaries; таблицы генерируются из "
        "JSON/CSV; 58 tests проходят; checkpoint fingerprint v2 включает content SHA-256; "
        "LaTeX проходит полный bib cycle; Word и PDF проверены постранично; "
        "публичный Git не содержит benchmark text, raw generations, model cache или "
        "внутренние рабочие материалы.",
    )
    paragraph(
        doc,
        "Перед отправкой остаются только авторские фактические поля: точное подразделение "
        "и адрес МГУ, funding, competing interests и окончательный author list. Квартиль "
        "проверяется непосредственно перед подачей с указанием базы, категории и года.",
    )

    doc.core_properties.title = (
        "Graph Rescue RAG: теория, расширенная проверка и пакет JIIS"
    )
    doc.core_properties.subject = (
        "Selective graph rescue for multi-hop information retrieval"
    )
    doc.core_properties.comments = (
        "Version 6.0. Includes global transfer, target calibration, an official-code "
        "baseline, clean latency, and the JIIS submission boundary."
    )
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
